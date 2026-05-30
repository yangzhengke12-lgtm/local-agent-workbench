"""将项目简历导出为 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_bold_para(text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    return p


def add_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style="Light List Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        t.cell(0, j).text = h
        for r in t.cell(0, j).paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            t.cell(i+1, j).text = text
    return t


# ── 标题 ──
title = doc.add_heading("Multi-Agent 层级协作系统", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("从单个 AI Agent 到 7 智能体工程团队的搭建过程").font.size = Pt(13)

doc.add_paragraph()

# ── 一句话 ──
add_bold_para("一句话：").add_run("从单个 DeepSeek API 对话机器人起步，逐步搭建成一个 7 智能体的工程团队模拟系统。")

# ── 技术栈 ──
doc.add_heading("技术栈", level=1)
doc.add_paragraph("Python · Anthropic SDK · ThreadPoolExecutor · JSON 驱动配置 · Subprocess 管理")

# ── 搭建过程 ──
doc.add_heading("搭建过程", level=1)

doc.add_heading("1. 起点 — 单 Agent", level=2)
doc.add_paragraph("从调用 DeepSeek API 的对话机器人起步，只有 2 个工具（获取时间、读文件）。跑通后发现了三个硬伤：上下文混乱、权限不可控、无法并行。")

doc.add_heading("2. 迭代 — 引入 Manager + Worker 分层", level=2)
doc.add_paragraph("Manager 负责拆解任务、指派员工、审核结果；5 个 Worker 各有独立系统提示和工具权限。用 workers.json 做零代码配置，ThreadPoolExecutor 实现并行调度。")

# Worker 表
add_table(
    ["Worker", "角色", "权限", "模型"],
    [
        ["亚历克斯", "高级开发", "读写执行搜索联网", "deepseek-v4-pro"],
        ["索菲亚", "代码审查", "只读搜索联网", "deepseek-v4-pro"],
        ["马库斯", "DevOps", "命令执行搜索时间", "deepseek-chat"],
        ["埃琳娜", "技术文档", "只读搜索联网", "deepseek-chat"],
        ["纳撒尼尔", "测试工程师", "读写执行搜索", "deepseek-v4-pro"],
    ],
)

doc.add_heading("3. 踩坑 — 三个独立排查修复的问题", level=2)

pitfalls = [
    ("Windows 编码崩溃", "find 命令在 Windows 下返回空 → 定位到 subprocess 用 GBK 解码非 GBK 字节导致 stdout 崩为 None → 改为 encoding='utf-8', errors='replace'"),
    ("Thinking 模式 API 400 错误", "模型返回的 thinking 块未被保留回传 → 在消息历史循环中补充 thinking 块处理，确保 thinking + signature 完整传递"),
    ("工具输出撑爆上下文", "单次 read_file 可能返回数万字符 → 自研截断保护，每个工具硬上限 6000 字符，搜索结果上限 50 条"),
]
for title_text, desc in pitfalls:
    p = doc.add_paragraph()
    p.add_run(f"• {title_text}：").bold = True
    p.add_run(desc)

doc.add_heading("4. 深度 — 从派活工具到工程组织", level=2)
features = [
    "任务看板：创建/分配/状态流转（todo → in_progress → done/failed）",
    "圆桌讨论：多 Worker 并行发言 + 两轮交叉回应 + 汇总共识",
    "共享知识库：经验沉淀 + 关键词检索，Worker 可直接查询",
    "三维绩效评分：正确性/完整性/质量各 1-5 分，持久化追踪趋势",
    "副经理制衡：发现 Manager 单点故障后引入，重大决策前独立复核",
    "GitHub 集成：create_pr / list_issues，与真实研发流程接轨",
    "文档转换：MarkItDown 集成，PDF/Word/PPT → Markdown",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

# ── 量化 ──
doc.add_heading("量化数据", level=1)
add_table(
    ["指标", "数值"],
    [
        ["智能体数量", "7（1 正经理 + 1 副经理 + 5 Worker）"],
        ["工具数量", "13 Worker 工具 + 13 管理工具"],
        ["代码量", "900+ 行 Python"],
        ["配置文件驱动", "100%（增删员工、调权限不改代码）"],
    ],
)

# ── 面试可讲的点 ──
doc.add_heading("面试可重点展开的点", level=1)
interview_points = [
    ("权限隔离的双保险", "API 层 tools 白名单 + 运行时二次校验，不是 prompt 软约束"),
    ("单点故障修复", "自研副经理机制：独立分析 + 明确反对 + 替代方案"),
    ("跨平台编码排查", "Windows GBK → subprocess stdout 崩为 None，独立定位根因并修复"),
    ("AI 辅助开发的边界", "架构设计、安全策略、踩坑修复是我做的；boilerplate 由 AI 生成"),
]
for title_text, desc in interview_points:
    p = doc.add_paragraph()
    p.add_run(f"• {title_text}：").bold = True
    p.add_run(desc)

# ── 保存 ──
output_path = "项目简历-Multi-Agent层级协作系统.docx"
doc.save(output_path)
print(f"已生成: {output_path}")
